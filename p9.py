import sys
from datetime import datetime
from PyQt6.QtWidgets import *
from PyQt6.uic import loadUiType
import docx
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd
from docx.enum.style import WD_STYLE_TYPE
ui1, _ = loadUiType("MAIN1.ui")


class Example(QWidget):

    def __init__(self):
        super().__init__()

        self.initUI()

    def initUI(self):

        self.setGeometry(300, 300, 300, 200)
        self.setWindowTitle('ComboBox')

        self.cb = QComboBox(self)
        self.cb.addItem('Option 1')
        self.cb.addItem('Option 2')
        self.cb.addItem('Option 3')
        self.cb.addItem('Other')
        self.cb.move(50, 50)
        self.cb.currentIndexChanged.connect(self.onActivated)

        self.show()

    def onActivated(self, index):

        text = self.cb.currentText()
        if text == 'Other':
            newItem, ok = QInputDialog.getText(
                self, 'New Item', 'Enter new item:')

            if ok:
                self.cb.addItem(newItem)
                self.cb.setCurrentIndex(self.cb.count() - 1)

                QMessageBox.information(
                    self, 'Success', f'"{newItem}" has been added to the ComboBox.')


class MainApp(QMainWindow, ui1):
    def __init__(self, parent=None):
        super(MainApp, self).__init__(parent)
        QMainWindow.__init__(self)
        self.setupUi(self)

        self.doc_name = None
        self.s_grade = None
        self.setupUi(self)
        self.doc = docx.Document()
        self.set_pageMargin()
        self.button_actions()
        self.acki = 97
        self.ackii = 97

    def button_actions(self):

        # -----------------------------------------------------------
        self.pushButton_Save.clicked.connect(self.save_doc)
        self.pushButton_HdrFtr.clicked.connect(self.set_HdrFtr)
        self.pushButton_PGraph.clicked.connect(self.Doc_paragraph)
        self.pushButton_Doc_Para_heading.clicked.connect(self.Doc_Para_heading)

        # --------------------------------------------------------
        self.pushButton_upload_csv.clicked.connect(self.upload_csv)
        self.pushButton_Table_generate.clicked.connect(self.tabular)
        self.pushButton_upload_img.clicked.connect(self.upload_img)
        self.pushButton_Image_generate.clicked.connect(self.img)
        self.pushButton_TableHeading.clicked.connect(self.HeadingTable)
        self.pushButton_LmUp.clicked.connect(self.upper_lm_port)
        self.pushButton_LmPragraph.clicked.connect(self.midle_lm_port1)
        # self.pushButton_LmSPragraph.clicked.connect(self.midle_lm_port2)
        self.pushButton_LmDn.clicked.connect(self.lower_lm_port)
        self.comboBox_Doc_SG.currentTextChanged.connect(self.doc_SG)
        # self.pushButton_DocSPGraph.clicked.connect(self.doc_subparagraph)
        # self.pushButton_LmReset.clicked.connect(self.reset1)
        # self.pushButton_DocSPGraph_Tag_reset.clicked.connect(self.reset2)
        self.pushButton_EmptyT.clicked.connect(self.empty_tbl)
        self.pushButton_Approval_table.clicked.connect(self.approvals_table)
        self.pushButton_revision_history.clicked.connect(self.rev_history)

    def reset1(self):
        self.acki = 97

    def reset2(self):
        self.ackii = 97

    def save_doc(self):
        doc = self.doc
        now = datetime.now().strftime("%S-%M-%H-%d-%m-%Y")
        doc.save(f"{now}.docx")

    def get_docRef(self):
        self.ref = self.lineEdit_DocRef.text()
        return self.ref

    def get_LmRef(self):
        ab = self.comboBox_LM_NO.currentText()
        return ab

    def get_LmDate(self):
        selected_date = self.calendarWidget_LM.selectedDate()
        return selected_date

    def get_To(self):
        ab = self.comboBox_LM_To.currentText()
        return ab

    def get_LmSGRD(self):
        ab = self.comboBox_LM_SG.currentText()
        return ab

    def get_LmForName(self):
        ab = self.comboBox_LM_NAME.currentText()
        return ab

    def get_LmForRank(self):
        ab = self.comboBox_LM_Rank.currentText()
        return ab

    def get_GrpOfLmFor(self):
        ab = self.comboBox_LM_GROUP.currentText()
        return ab

    def get_tellOfLmFor(self):
        ab = self.comboBox_LM_TELL.currentText()
        return ab

    def get_LmSubject(self):
        sb = self.lineEdit_lmSubject.text()
        return sb

    def Lm_paragraph_text1(self):
        para = self.plainTextEdit_LmParagraph.toPlainText().strip()
        return para

    def Lm_paragraph_text2(self):
        para = self.plainTextEdit_LmParagraph_2.toPlainText().strip()
        return para

    def Lm_paragraph_text3(self):
        para = self.plainTextEdit_LmParagraph_3.toPlainText().strip()
        return para

    def Lm_paragraph_text4(self):
        para = self.plainTextEdit_LmParagraph_4.toPlainText().strip()
        return para

    def Lm_Sub_paragraph_text(self):
        para = self.plainTextEdit_LmSParagraph.toPlainText().strip()
        return para

    def get_docRev(self):
        self.rev = self.lineEdit_DocRev.text()
        return self.rev

    def get_docDate(self):
        selected_date = self.calendarWidget_Doc.selectedDate()
        return selected_date

    def empty_tbl_rows(self):
        row = self.spinBox_TEmpty_Row.text()
        return row

    def empty_tbl_cols(self):
        col = self.spinBox_TEmpty_Col.text()
        return col

    def empty_tbl(self):
        doc = self.doc
        n_rows = self.empty_tbl_rows()
        n_cols = self.empty_tbl_cols()
        width = 7
        n_rows, n_cols, width = map(int, [n_rows, n_cols, width])
        table = doc.add_table(rows=n_rows, cols=n_cols)

        table.style = "Table Grid"
        for row in table.rows:
            row.height = Inches(0.5)

    def add_table(self, df):
        doc = self.doc
        table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])
        table.style = "Table Grid"
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for row in table.rows:
            row.height = Inches(0.5)
        header = table.rows[0].cells
        for i in range(df.shape[1]):
            header[i].text = df.columns[i]
            if i == 0:
                header[i].paragraphs[0].runs[0].font.bold = True
            header[i].paragraphs[0].runs[0].font.size = Pt(12)
            header[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            header[i].paragraphs[0].runs[0].font.name = 'Arial'
        for i in range(df.shape[0]):
            row = table.rows[i+1].cells
            for j in range(df.shape[1]):
                cell = row[j].add_paragraph(str(df.values[i, j]))
                cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.runs[0].font.size = Pt(12)
                cell.runs[0].font.name = 'Arial'

                if j == 0:
                    cell.runs[0].font.bold = True

    def upload_csv(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Open CSV File", ".", "CSV Files (*.csv)")
        if file_path:
            self.lineEdit_csv_FilePath.setText(file_path)

    def tbl_heading(self):
        t_head = self.lineEdit_table_heading.text()
        t_head = t_head.upper()
        return t_head

    def tabular(self):
        t = self.tbl_heading()
        self.i_heading_handler(t)
        self.draw_table()

    def draw_table(self):
        # Read CSV file using pandas
        file_path = self.lineEdit_csv_FilePath.text()
        if file_path:
            df = pd.read_csv(file_path)
            self.add_table(df)

    def get_document_name(self):
        self.doc_name = self.lineEdit_DocName.text()
        return self.doc_name

    def LmHeader(self, headerText):
        doc = self.doc
        header = doc.sections[0].header
        paragraph = header.paragraphs[0]
        paragraph.add_run(headerText)
        # set the paragraph properties
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        paragraph.style.font.name = 'Arial'
        paragraph.style.font.size = docx.shared.Pt(12)
        paragraph.style.font.bold = False

    def LmFooter(self, footerText):
        doc = self.doc
        footer = doc.sections[0].footer
        paragraph = footer.paragraphs[0]
        paragraph.add_run(footerText)
        # set the paragraph properties
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        paragraph.style.font.name = 'Arial'
        paragraph.style.font.size = docx.shared.Pt(12)
        paragraph.style.font.bold = False

    def LmHdrFtr(self, grade):
        self.LmHeader(grade)
        self.LmFooter(grade)

    def para_space_handler(self):
        doc = self.doc
        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.line_spacing = 1.0

        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.space_after = 0

    def upper_lm_port(self):
        sg = self.get_LmSGRD()
        self.LmHdrFtr(sg)
        to = self.get_To()
        To = to.upper()
        sub = self.get_LmSubject()
        SUB = sub.upper()
        doc = self.doc
        doc.add_paragraph("")
        doc.add_paragraph("")
        org_name = "Avionics Production Factory".upper()
        title = doc.add_paragraph(org_name)
        title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        title.style.font.bold = True
        title.style.font.name = "Arial"
        title.style.font.size = docx.shared.Pt(12)
        run = title.runs[0]
        run.font.bold = True

        sub_title = "(DDD)"
        stitle = doc.add_paragraph(sub_title)
        stitle.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        stitle.style.font.bold = False
        stitle.style.font.name = "Arial"
        stitle.style.font.size = docx.shared.Pt(12)

        doc.add_paragraph("")
        doc.add_paragraph("")

        ref_to = doc.add_paragraph(To)
        ref_to.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT

        doc.add_paragraph("")
        doc.add_paragraph("")

        if SUB == '':
            sub = "Subject".upper()
        else:
            sub = SUB.upper()
        start = doc.add_paragraph(sub)
        start.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        start.style.font.name = "Arial"
        run = start.runs[0]
        run.font.bold = True
        run.font.size = Pt(14)
        doc.add_paragraph("")
        self.para_space_handler()

    def midle_lm_port1(self):
        doc = self.doc
        doc.add_paragraph("")
        your_para1 = self.Lm_paragraph_text1()
        your_para2 = self.Lm_paragraph_text2()
        your_para3 = self.Lm_paragraph_text3()
        your_para4 = self.Lm_paragraph_text4()

        pragraphs = [your_para1, your_para2, your_para3, your_para4]
        for your_para in pragraphs:
            if your_para == '':
                continue  # skip empty paragraphs
            else:
                paragraph = your_para

            paragraph = '\t' + paragraph
            paragraph = doc.add_paragraph(paragraph)
            paragraph.style = 'List Number'
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(12)
            font.bold = False

        self.para_space_handler()

    def midle_lm_port2(self):
        doc = self.doc
        ascii_val = self.acki
        char_val = chr(ascii_val)
        doc.add_paragraph("")
        paragraph = self.Lm_Sub_paragraph_text()
        paragraph = '\t'+f'({char_val})' + paragraph
        paragraph = doc.add_paragraph(paragraph)
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)
        font.bold = False
        self.para_space_handler()
        self.acki += 1

    def doc_subpara_text(self):
        doc_subPara = self.textEdit_SparaGraph_text.toPlainText().strip()
        return doc_subPara

    def doc_subparagraph(self):
        doc = self.doc
        ascii_val1 = self.ackii
        char_val = chr(ascii_val1)
        doc.add_paragraph("")
        paragraph = self.doc_subpara_text()
        paragraph = '\t' + f'({char_val})' + paragraph
        paragraph = doc.add_paragraph(paragraph)
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)
        font.bold = False
        self.para_space_handler()
        self.ackii += 1

    def stamp_maker(self, NAME, RANK, GROUP, TEL):
        doc = self.doc

        # determine maximum length of each field
        max_name_len = 20
        max_rank_len = 20
        max_group_len = 20
        max_tel_len = 20

        # add a left indent of 7 steps
        left_indent = 24 * docx.shared.Pt(12)

        # left justify the name field with left indent
        name = doc.add_paragraph()
        name.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
        name.paragraph_format.left_indent = left_indent
        name.add_run(f"{NAME}".ljust(max_name_len))
        name.style.font.name = "Arial"
        name.style.font.size = Pt(12)
        name.style.font.bold = True

        # left justify the rank field with left indent
        rank = doc.add_paragraph()
        rank.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
        rank.paragraph_format.left_indent = left_indent
        rank.add_run(f"{RANK}".ljust(max_rank_len))
        rank.style.font.name = "Arial"
        rank.style.font.size = Pt(12)
        rank.style.font.bold = False

        # left justify the group field with left indent
        group = doc.add_paragraph()
        group.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
        group.paragraph_format.left_indent = left_indent
        group.add_run(f"{GROUP}".ljust(max_group_len))
        group.style.font.name = "Arial"
        group.style.font.size = Pt(12)
        group.style.font.bold = False

        # left justify the tel field with left indent
        tel = doc.add_paragraph()
        tel.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
        tel.paragraph_format.left_indent = left_indent
        tel.add_run(f"Tel Ext: {TEL}".ljust(max_tel_len))
        tel.style.font.name = "Arial"
        tel.style.font.size = Pt(12)
        tel.style.font.bold = False

    def lower_lm_port(self):
        doc = self.doc
        for i in range(5):
            doc.add_paragraph("")
        lm_no = self.get_LmRef()
        date = self.get_LmDate()
        nam = self.get_LmForName()
        rnk = self.get_LmForRank()
        grp = self.get_GrpOfLmFor()
        tel = self.get_tellOfLmFor()
        formatted_date = date.toString("dd MMMM, yyyy")
        LM_REFRENCE = lm_no.upper()
        # DATE = date.upper()
        NAME = nam.upper()
        RANK = rnk.upper()
        GROUP = grp.upper()
        TEL = tel.upper()
        self.stamp_maker(NAME, RANK, GROUP, TEL)
        doc.add_paragraph("")
        lm_ref = doc.add_paragraph()
        lm_ref.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
        lm_ref.style.font.bold = False
        lm_ref.style.font.name = "Arial"

        lm_ref.add_run("LM No  ")
        lm_ref.add_run(LM_REFRENCE)
        lm_ref.add_run(" ")
        lm_ref.add_run("   dated    ")
        lm_ref.add_run(formatted_date)

        self.para_space_handler()

    def doc_SG(self):
        ab = self.comboBox_Doc_SG.currentText()
        return ab

    def set_HdrFtr(self):
        doc_name = self.get_document_name()
        doc_sgrade = self.doc_SG()
        doc_ref_no = self.get_docRef()
        doc_rev_no = self.get_docRev()
        doc_date = self.get_docDate()

        self.set_header(doc_name, doc_sgrade)
        self.set_footer(doc_sgrade, doc_ref_no, doc_rev_no, doc_date)

    def set_pageMargin(self):
        doc = self.doc
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(1.5)
            section.right_margin = Inches(0.5)

    def get_table_sHd(self):
        s_hd = self.lineEdit_TableHeadS.text()
        s_hd = s_hd.upper()
        return s_hd

    def get_table_lHd(self):
        l_hd = self.lineEdit_TableHeadL.text()
        l_hd = l_hd.upper()
        return l_hd

    def HeadingTable(self):
        s = self.get_table_sHd()
        l = self.get_table_lHd()
        print(s)
        print(l)
        self.table_heading(s, l)

    def table_heading(self, s_hd, l_hd):
        doc = self.doc
        doc.add_paragraph("")

        table = doc.add_table(rows=2, cols=1)
        table.style = 'Table Grid'

        table.cell(0, 0).text = s_hd
        table.cell(1, 0).text = l_hd

        run = table.cell(0, 0).paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(12)

        run = table.cell(1, 0).paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(14)

        table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(
            0, 0).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.cell(1, 0).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        table.cell(
            1, 0).paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        table.rows[0].height = Inches(0.5)
        table.rows[1].height = Inches(0.5)

    def set_header(self, your_docNm, your_sGrd):
        doc = self.doc
        header = doc.sections[0].header
        if your_docNm == '':
            ui_hdr_doc_nam = "my new document for test"
        else:
            ui_hdr_doc_nam = your_docNm
        if your_sGrd == '':
            ui_hdr_s_grd = "confidential"
        else:
            ui_hdr_s_grd = your_sGrd
        ui_hdr_doc_nam = ui_hdr_doc_nam.upper()
        # ui_hdr_s_grd = ui_hdr_s_grd.upper()
        print(ui_hdr_doc_nam)
        print(ui_hdr_s_grd)
        table = header.add_table(rows=1, cols=3, width=Inches(8.01))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Resize the table to fit the header
        # for all rows
        for row in table.rows:
            for cell in row.cells:
                cell.width = docx.shared.Inches(2.67)
        # first cell
        first_cell = table.cell(0, 0)
        first_cell.text = ui_hdr_doc_nam
        self.set_cell_font(first_cell, 12)

        # second cell
        second_cell = table.cell(0, 1)
        second_cell.text = ui_hdr_s_grd
        self.set_cell_font(second_cell, 12)
        # third cell
        third_cell = table.cell(0, 2)
        run = third_cell.paragraphs[0].add_run()
        picture = run.add_picture("pic.jpg")
        picture.width = docx.shared.Inches(0.5)
        picture.height = docx.shared.Inches(0.25)
        third_cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        third_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # header from top
        doc.sections[0].header_distance = Inches(0.5)
        # header row height
        for row in table.rows:
            row.height = Inches(0.5)

    def set_footer(self, sGrd, dRef, dRev, dDate):
        doc = self.doc
        # Set the footer of the document
        section = doc.sections[-1]
        footer = section.footer
        # Create a table with one row and three columns
        table = footer.add_table(rows=2, cols=4, width=Inches(8))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        # Add content to the cells
        cell1 = table.cell(0, 0)
        cell1.text = "Document Ref".upper()
        cell1.width = Inches(2.5)
        self.set_cell_font(cell1, 12)

        cell2 = table.cell(0, 1)
        cell2.text = "Rev No".upper()
        cell2.width = Inches(1.5)
        self.set_cell_font(cell2, 12)

        cell3 = table.cell(0, 2)
        cell3.text = "DATE"
        cell3.width = Inches(2)
        self.set_cell_font(cell3, 12)

        cell4 = table.cell(0, 3)
        cell4.text = 'Page'
        cell4.width = Inches(2.5)
        self.set_cell_font(cell4, 12)

        cell5 = table.cell(1, 0)
        if dRef == '':
            doc_Ref = 'Enter Doc Ref'
        else:
            doc_Ref = dRef
        doc_Ref = doc_Ref.upper()
        print(f"Ref No In Footer : {doc_Ref}")
        cell5.text = doc_Ref
        cell5.width = Inches(2)
        self.set_cell_font(cell5, 12)

        cell6 = table.cell(1, 1)
        if dRev == '':
            doc_Rev = 'Enter Rev No'
        else:
            doc_Rev = dRev
        doc_Rev = doc_Rev.upper()
        print(f"Rev No In Footer : {doc_Rev}")
        cell6.text = doc_Rev
        cell6.width = Inches(1.5)
        self.set_cell_font(cell6, 12)

        cell7 = table.cell(1, 2)
        if dDate == '':
            doc_Date = 'Enter Date'
        else:
            doc_Date = dDate
        formatted_date = doc_Date.toString("dd MMMM, yyyy")
        # doc_Date = doc_Date.upper()
        print(f"Date In Footer : {doc_Date}")
        cell7.text = formatted_date
        cell7.width = Inches(2)
        self.set_cell_font(cell7, 12)

        cell8 = table.cell(1, 3)
        run = cell8.paragraphs[0].add_run()
        field = OxmlElement('w:fldSimple')
        field.set(qn('w:instr'), 'PAGE')
        run._r.append(field)
        run.add_text(' of ')
        field = OxmlElement('w:fldSimple')
        field.set(qn('w:instr'), 'NUMPAGES')
        run._r.append(field)
        self.set_cell_font(cell8, 12)
        paragraph = footer.add_paragraph()
        if sGrd == '':
            p = "confidential"
        else:
            p = sGrd
        p = p.upper()
        print(f"Security Grade In Footer : {p}")
        footer_run = paragraph.add_run(p)
        footer_run.bold = False
        footer_run.font.size = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Inches(0.2)
        paragraph_format.space_after = Inches(0.2)
        # header from top
        doc.sections[0].footer_distance = Inches(0.5)
        # footer row height
        for row in table.rows:
            row.height = Inches(0.5)

    def Doc_Para_heading(self):
        heading = self.lineEdit_paragraph_heading.text()
        heading = heading.upper()
        self.p_heading_handler(heading)
        return heading

    def get_Img_heading(self):
        heading = self.lineEdit_img_heading.text()
        heading = heading.upper()

    def p_heading_handler(self, heading):
        p_head = heading
        doc = self.doc
        doc.add_paragraph("")
        heading = doc.add_paragraph(p_head)
        font = heading.style.font
        font.name = 'Arial'
        font.size = Pt(12)
        for run in heading.runs:
            run.font.bold = True
        heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    def i_heading_handler(self, heading):
        p_head = heading
        doc = self.doc
        doc.add_paragraph("")
        heading = doc.add_paragraph(p_head)
        font = heading.style.font
        font.name = 'Arial'
        font.size = Pt(12)
        for run in heading.runs:
            run.font.bold = False
        heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

    def Doc_para_text(self):
        P_text = self.textEdit_Docs_paraGraph_text.toPlainText()
        return P_text

    def img(self):
        img_heading = self.get_Img_heading()
        self.i_heading_handler(img_heading)
        self.draw_img()

    def Lm_img(self):
        self.draw_img()

    def para(self):
        your_heading = self.get_P_heading()
        self.p_heading_handler(your_heading)
        self.add_paragraph1()

    def Doc_paragraph(self):
        doc = self.doc

        your_para = self.Doc_para_text()
        if your_para == '':
            paragraph = "This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. This is Dummy Paragraph. "
        else:
            paragraph = your_para

        paragraph = '\t' + paragraph
        paragraph = doc.add_paragraph(paragraph)
        paragraph.style = 'List Number'
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)
        font.bold = False
        self.para_space_handler()

    def set_cell_font(self, cell, size):
        # cell.paragraphs[0].runs[0].font.size = Pt(size)
        cell.paragraphs[0].style.font.size = Pt(size)
        cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].style.font.bold = False
        cell.paragraphs[0].style.font.name = 'Arial'
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def upload_img(self):
        self.file_path, _ = QFileDialog.getOpenFileName(
            self, 'Select Image File', '', 'JPEG Files (*.jpg);;JPG Files (*.jpeg);;PNG Files (*.png)')
        if self.file_path:
            self.lineEdit_img_FilePath.setText(self.file_path)

    def draw_img(self):
        # Read CSV file using pandas
        file_path = self.lineEdit_img_FilePath.text()
        if file_path:
            self.insert_img(file_path)

    def insert_img(self, path_img):
        doc = self.doc
        doc.add_paragraph()
        doc.add_picture(path_img, width=Inches(4))
        inline_shape = doc.inline_shapes[-1]
        width, height = inline_shape.width, inline_shape.height
        aspect_ratio = int(height / width)
        new_width = Inches(2)
        new_height = aspect_ratio * new_width
        inline_shape.width = new_width
        inline_shape.height = new_height
        paragraph = doc.paragraphs[-1]
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    def insert_LmImg(self, pic_path):
        doc = self.doc
        doc.add_paragraph()
        doc.add_picture(pic_path, width=Inches(4))
        inline_shape = doc.inline_shapes[-1]
        width, height = inline_shape.width, inline_shape.height
        aspect_ratio = int(height / width)
        new_width = Inches(2)
        new_height = aspect_ratio * new_width
        inline_shape.width = new_width
        inline_shape.height = new_height
        paragraph = doc.paragraphs[-1]
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    def rev_history(self):
        doc = self.doc
        P = doc.add_paragraph("")
        p_head11 = "REVISION HISTORY"
        heading11 = doc.add_paragraph(p_head11)
        font = heading11.style.font
        font.name = 'Arial'
        font.size = Pt(12)
        for run in heading11.runs:
            run.font.bold = False
        heading11.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        table = doc.add_table(rows=2, cols=7)
        table.style = 'Table Grid'

        # Add headings to the table cells
        headings = ["Date", "Rev", "Doc Version No", "No. of Changes",
                    "Page No", "Applicability", "Description"]
        for j in range(7):
            cell = table.cell(0, j)
            cell.text = headings[j]
            # Bold the text in the first row
            cell.paragraphs[0].runs[0].bold = True

        for row in table.rows:
            row.height = Inches(0.5)
        for row in table.rows:
            for cell in row.cells:
                cell.width = docx.shared.Inches(1)

        # Set horizontal and vertical alignment to justify
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER

    def approvals_table(self):

        # Create a new Word document
        doc = self.doc
        p_head1 = "APPROVALS"
        heading1 = doc.add_paragraph(p_head1)
        font = heading1.style.font
        font.name = 'Arial'
        font.size = Pt(12)
        for run in heading1.runs:
            run.font.bold = False
        heading1.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        # Add a table with 10 rows and 5 columns
        table = doc.add_table(rows=10, cols=5)
        table.style = 'Table Grid'

        # Set the height of each cell to 0.5 inches
        for row in table.rows:
            for cell in row.cells:
                cell.height = docx.shared.Inches(0.5)

        # Merge cells in the first column
        cell_0_0 = table.cell(0, 0)
        cell_1_0 = table.cell(1, 0)
        cell_2_0 = table.cell(2, 0)
        cell_3_0 = table.cell(3, 0)
        cell_4_0 = table.cell(4, 0)
        cell_0_0.merge(cell_1_0)
        cell_0_0.merge(cell_2_0)
        cell_0_0.merge(cell_3_0)
        cell_0_0.merge(cell_4_0)

        # Merge cells in the second column
        cell_1_1 = table.cell(1, 1)
        cell_2_1 = table.cell(2, 1)
        cell_2_1.merge(cell_1_1)

        # Merge cells in the second column
        cell_3_1 = table.cell(3, 1)
        cell_4_1 = table.cell(4, 1)
        cell_4_1.merge(cell_3_1)

        # Merge cells in the second column
        cell_5_1 = table.cell(5, 1)
        cell_6_1 = table.cell(6, 1)
        cell_6_1.merge(cell_5_1)

        # Merge cells in the second column
        cell_7_1 = table.cell(7, 1)
        cell_8_1 = table.cell(8, 1)
        cell_8_1.merge(cell_7_1)

        # Merge cells in row 6 and 7 of column 1
        cell_5_0 = table.cell(5, 0)
        cell_6_0 = table.cell(6, 0)
        cell_5_0.merge(cell_6_0)

        # Merge cells in row 8 and 9 of column 1
        cell_7_0 = table.cell(7, 0)
        cell_8_0 = table.cell(8, 0)
        cell_7_0.merge(cell_8_0)

        # Add headings to the table cells
        headings = ["Prepared by", "Concerned Sections",
                    "Rank/Name", "Signature", "Date"]
        for j in range(5):
            cell = table.cell(0, j)
            cell.text = headings[j]

        cell_5_0.text = "Verified by"
        cell_7_0.text = "Issued by"
        cell_9_0 = table.cell(9, 0)
        cell_9_0.text = "Approved by"

        for row in table.rows:
            row.height = Inches(0.5)

        # Set horizontal and vertical alignment to justify
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
                cell.vertical_alignment = docx.enum.table.WD_ALIGN_VERTICAL.CENTER


def main():
    app = QApplication(sys.argv)
    window = MainApp()
    window.show()
    app.exec()


if __name__ == "__main__":
    main()
